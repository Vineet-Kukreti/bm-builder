"""
BM Builder — Orchestration engine core (engine.core).

Copyright 2026 Vineet Kukreti (Bespoke Mind AI) — https://bespokemind.ai
Licensed under the Apache License, Version 2.0. See LICENSE and NOTICE.

Responsibilities:
  * Route each agent's prompt to the AI the operator chose for it (per-agent in
    Settings). Providers:
        - Claude subscription : Claude via Claude Code on the logged-in plan ($0).
        - "anthropic"         : Claude models via the Anthropic API (metered).
        - "openai"            : OpenAI / OpenAI-compatible endpoints (Groq, OpenRouter…).
  * Hold the standing personas for the executive team (CEO / CTO / CMO) and the
    delivery team (Developer / QA / Designer).
  * Manage a self-contained, per-project workspace on local disk:
        workspace_builds/<project>/
            prd.md  plan.md
            src/                   generated source artifacts
            history/changelog.md   append-only log of every change/decision
            bugs.md                bug knowledge base (QA re-reads it each pass)
"""

import os
import re
import json
import base64
import shutil
import tempfile
import threading
import subprocess
from io import BytesIO
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

# Repo root (parent of this engine/ package) so .env / settings.json / lessons.json stay at the
# project root regardless of where this module physically lives.
_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Load variables from a local .env file (e.g. ANTHROPIC_API_KEY) if present.
# Does not override variables already set in the real environment.
try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(_REPO_ROOT, ".env"))
except ImportError:
    pass

# --------------------------------------------------------------------------
# Providers (lazily constructed so the app still loads if one is unavailable)
# --------------------------------------------------------------------------
# Cloud OpenAI endpoint, used when the operator picks an "OpenAI" provider in Settings
# (a custom "OpenAI-compatible" provider supplies its own base URL).
_OPENAI_CLOUD_BASE = "https://api.openai.com/v1"

# The global default agent model, set from Settings (see apply_settings). Provider is one
# of: "claude_subscription" (the $0 default — agents run on the logged-in Claude plan),
# "anthropic" (metered Claude API), "openai" (cloud), "openai_compatible" (custom base URL).
# This drives ALL call_agent agents; the autonomous Claude Code BUILD always uses the
# subscription regardless (OpenAI can't drive the `claude` CLI).
_DEFAULT_MODEL = {"provider": "claude_subscription", "model": "", "base_url": ""}

# Per-agent provider overrides {ROLE: provider}, set from Settings. A role not listed here
# falls back to _DEFAULT_MODEL["provider"] (so an empty map == "all agents use the default").
_AGENT_PROVIDERS = {}

_oai_clients = {}
_anthropic_client = None


def openai_client(base_url):
    """Cached OpenAI-compatible client for a base URL. Uses OPENAI_API_KEY from the
    environment / Settings (a placeholder key is sent if none is set)."""
    if base_url not in _oai_clients:
        from openai import OpenAI
        key = os.environ.get("OPENAI_API_KEY") or "local"
        _oai_clients[base_url] = OpenAI(base_url=base_url, api_key=key)
    return _oai_clients[base_url]


def reset_clients():
    """Drop cached provider clients so an API-key or base-URL change in Settings takes
    effect immediately, without restarting the app."""
    global _anthropic_client
    _anthropic_client = None
    _oai_clients.clear()


def anthropic_client():
    """Anthropic client, or None if the package/API key is unavailable."""
    global _anthropic_client
    if _anthropic_client is None:
        if not os.environ.get("ANTHROPIC_API_KEY"):
            return None
        try:
            import anthropic
        except ImportError:
            return None
        _anthropic_client = anthropic.Anthropic()
    return _anthropic_client


def engine_status():
    """Report which back-ends are ready, for display in the UI."""
    try:
        import anthropic  # noqa: F401
        pkg = True
    except ImportError:
        pkg = False
    return {
        "anthropic_pkg": pkg,
        "anthropic_key": bool(os.environ.get("ANTHROPIC_API_KEY")),
    }


# --------------------------------------------------------------------------
# Model registry — the menu the CEO picks from when staffing a project.
# Each entry: provider, whether it can see images, and a one-line rationale
# so the CEO can choose intelligently per project.
# --------------------------------------------------------------------------
MODEL_REGISTRY = {
    # Claude (Anthropic) — vision-capable. On the Claude subscription these run at $0 via
    # Claude Code; on the metered API they bill at the prices below.
    "claude-opus-4-8": {
        "provider": "anthropic", "vision": True,
        "desc": "Most capable. Best for hard architecture, strategy and long-horizon planning.",
    },
    "claude-sonnet-4-6": {
        "provider": "anthropic", "vision": True,
        "desc": "Strong balance of speed and intelligence. A good default for most agents.",
    },
    "claude-haiku-4-5": {
        "provider": "anthropic", "vision": True,
        "desc": "Fast and economical. Good for simple, high-volume or latency-sensitive tasks.",
    },
    # OpenAI — vision-capable, metered (used when an agent is set to the OpenAI provider).
    "gpt-4o": {
        "provider": "openai", "vision": True,
        "desc": "OpenAI flagship multimodal model. Strong general reasoning; vision-capable.",
    },
    "gpt-4o-mini": {
        "provider": "openai", "vision": True,
        "desc": "Fast, economical OpenAI model. Good for high-volume or simpler tasks.",
    },
}

# Which API family a configured provider belongs to (drives model choice + the call backend).
_PROVIDER_FAMILY = {"claude_subscription": "anthropic", "anthropic": "anthropic",
                    "openai": "openai", "openai_compatible": "openai"}


def provider_family(prov):
    """'anthropic' (Claude — subscription or API) or 'openai' (OpenAI / OpenAI-compatible)."""
    return _PROVIDER_FAMILY.get(prov, "anthropic")

# Anthropic list prices, USD per 1M tokens (input, output).
PRICING = {
    "claude-fable-5": (10.0, 50.0),
    "claude-opus-4-8": (5.0, 25.0),
    "claude-opus-4-7": (5.0, 25.0),
    "claude-opus-4-6": (5.0, 25.0),
    "claude-sonnet-4-6": (3.0, 15.0),
    "claude-haiku-4-5": (1.0, 5.0),
}

# The CEO is fixed (decision-maker). The CTO is pinned to the latest flagship so
# its architecture guidelines and code review always run on a top model.
CEO_MODEL = "claude-opus-4-8"
CTO_MODEL = "claude-opus-4-8"

# Roles the CEO assigns a model to (CEO and CTO are pinned, so excluded here).
ASSIGNABLE_ROLES = ["CMO", "DEVELOPER", "QA", "DESIGNER", "PM"]

# The cloud "thinking" agents whose AI provider the operator can pick individually in Settings.
# (DEVELOPER/DESIGNER are delivery roles: the autonomous build always uses the subscription and
# visual review needs a vision model, so they aren't operator-routable here.)
ROUTABLE_ROLES = ["CEO", "CTO", "CMO", "PM", "QA", "SKEPTIC"]

# Used before the CEO has assigned, or when an assignment is invalid. All Claude:
# Opus for the heavy reasoning/alignment roles, Sonnet as the balanced default,
# and the DESIGNER on a vision-capable Claude model.
DEFAULT_ASSIGNMENT = {
    "CMO": "claude-sonnet-4-6",
    "DEVELOPER": "claude-sonnet-4-6",
    "QA": "claude-sonnet-4-6",
    "DESIGNER": "claude-opus-4-8",
    "PM": "claude-sonnet-4-6",
}

CORE_TEAM = ["CEO", "CTO", "CMO"]

PERSONAS = {
    "CEO": (
        "You are the CEO of an AI solutions studio. You own business viability, "
        "scope, priorities, risk and hiring. You proactively surface blindspots that "
        "neither the operator nor the client mentioned, challenge weak assumptions, "
        "and keep the project commercially sound. Be concise and decisive."
    ),
    "CTO": (
        "You are the CTO of an AI solutions studio. You own technical architecture, "
        "feasibility, tech-stack selection, scalability, security and engineering risk. "
        "You proactively flag technical constraints, integration pitfalls and hidden "
        "complexity the client may not have considered. Be concrete and pragmatic."
    ),
    "CMO": (
        "You are the CMO of an AI solutions studio. You own the market, end users, "
        "positioning, UX and growth. You proactively suggest user-facing features and "
        "differentiators the client and operator may have overlooked, and you guard the "
        "product against poor usability. Be insightful and user-focused."
    ),
    "DEVELOPER": (
        "You are a master software engineer. You implement exactly what the frozen plan "
        "and PRD describe, writing clean, runnable code. Output raw code blocks, each "
        "clearly labelled with its filename."
    ),
    "QA": (
        "You are a strict QA engineer. You review BOTH backend/code logic AND the "
        "front-end. You check whether previously-logged bugs have regressed. You report "
        "findings as a structured list, each item tagged [CODE] or [FRONTEND] with a severity."
    ),
    "DESIGNER": (
        "You are a UI/UX designer with a sharp eye for visual quality. You review "
        "front-end screenshots for layout, hierarchy, spacing, accessibility, "
        "responsiveness and brand consistency, and propose concrete improvements."
    ),
    "PM": (
        "You are the Project Manager / delivery lead. You coordinate the Developer, QA and the "
        "rest of the team, keep the build on track and aligned to the PRD and delivery plan, and "
        "own final acceptance: after QA passes you review the delivered code and outputs against "
        "the PRD and the plan and flag any gaps or misalignments."
    ),
    "SKEPTIC": (
        "You are a ruthless technical reviewer / red-team. You stress-test plans and specs and "
        "hunt for missing requirements, contradictions, ambiguity, infeasibilities, security and "
        "privacy gaps, untestable acceptance criteria, and risky assumptions. Be specific and "
        "constructive — every issue gets a concrete recommendation."
    ),
}


def persona_for(role):
    return PERSONAS.get(role.upper(), "")


# --------------------------------------------------------------------------
# Model assignment helpers
# --------------------------------------------------------------------------
def registry_summary():
    """Human-readable model menu to hand the CEO when it assigns the roster."""
    lines = []
    for name, meta in MODEL_REGISTRY.items():
        vis = "vision" if meta["vision"] else "text-only"
        lines.append(f"- {name} [{meta['provider']}, {vis}]: {meta['desc']}")
    return "\n".join(lines)


def _fallback_vision_model():
    """A vision-capable model for image calls (visual review). Claude Opus is vision-capable;
    it needs an Anthropic API key (vision can't run through the Claude Code subscription)."""
    return "claude-opus-4-8"


def _default_model_for_family(fam):
    """A sensible default model for a provider family."""
    return "gpt-4o" if fam == "openai" else "claude-sonnet-4-6"


def validate_assignment(raw, fams=None):
    """Clean a CEO-proposed model assignment: drop unknown models, fill gaps with defaults,
    keep each role within its REQUIRED provider family (`fams[role]`, if given), and guarantee
    the DESIGNER (visual reviewer) is on a vision-capable model."""
    raw = raw or {}
    fams = fams or {}
    clean = {}
    for role in ASSIGNABLE_ROLES:
        choice = (raw.get(role) or "").strip()
        fam = fams.get(role)
        ok = choice in MODEL_REGISTRY and (fam is None or MODEL_REGISTRY[choice]["provider"] == fam)
        if not ok:
            choice = _default_model_for_family(fam) if fam else DEFAULT_ASSIGNMENT[role]
        clean[role] = choice
    # The designer performs visual review — it must be able to see images.
    if not MODEL_REGISTRY[clean["DESIGNER"]]["vision"]:
        clean["DESIGNER"] = _fallback_vision_model()
    return clean


def model_for(role, assignment=None):
    """Resolve which model a role runs on for the current project."""
    role = role.upper()
    if role == "CEO":
        return CEO_MODEL                       # orchestrator — always fixed
    if assignment and assignment.get(role):
        return assignment[role]                # explicit per-build choice wins (incl. CTO)
    if role == "CTO":
        return CTO_MODEL                       # CTO default = pinned latest flagship
    return DEFAULT_ASSIGNMENT.get(role, CEO_MODEL)


# --------------------------------------------------------------------------
# Model invocation
# --------------------------------------------------------------------------
def _openai_user_content(user_prompt, images):
    """Build the user content for OpenAI-compatible servers (image_url blocks)."""
    if not images:
        return user_prompt or ""
    content = [{"type": "text", "text": user_prompt or ""}]
    for img in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['media_type']};base64,{img['b64']}"},
        })
    return content


def _anthropic_user_content(user_prompt, images):
    """Build the user content for Anthropic (base64 image source blocks)."""
    if not images:
        return user_prompt or ""
    content = []
    for img in images:
        content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": img["media_type"], "data": img["b64"]},
        })
    content.append({"type": "text", "text": user_prompt or ""})
    return content


# Substrings that mark an OpenAI(-compatible) model as multimodal, so vision calls can
# use the operator's chosen model when it can actually see images (e.g. gpt-4o); otherwise
# they fall back to a vision-capable model so visual review never silently breaks.
_VISION_HINTS = ("4o", "4.1", "4-turbo", "vision", "omni", "o1", "o3", "gpt-5",
                 "-vl", "pixtral", "llava", "gemini")


def _openai_is_vision(model_id):
    m = (model_id or "").lower()
    return any(h in m for h in _VISION_HINTS)


def _vision_fallback():
    """(model_id, provider, base_url) for a vision-capable model when the chosen default
    can't see images — Anthropic if a key is configured, else a local vision model."""
    vmid = _fallback_vision_model()
    meta = MODEL_REGISTRY.get(vmid, {"provider": "anthropic"})
    return (vmid, meta["provider"], "")


def role_provider(role):
    """The AI provider configured for one agent: its per-agent Setting if set, else the global
    default. One of PROVIDER_CHOICES ('claude_subscription' | 'anthropic' | 'openai' |
    'openai_compatible')."""
    return (_AGENT_PROVIDERS.get((role or "").upper())
            or (_DEFAULT_MODEL or {}).get("provider", "claude_subscription"))


def _resolve_call(role, model, has_images):
    """Resolve (model_id, provider, base_url) for a call.

    The PROVIDER follows the operator's per-agent Settings (`role_provider`). The MODEL is the
    CEO's assignment for the role (passed as `model`/looked up via `model_for`) when it belongs to
    that provider's family, else a sensible default for the family. For 'openai_compatible' the
    custom model + base URL come from Settings. The autonomous Claude Code build does NOT pass
    through here; it always runs on the subscription."""
    dm = _DEFAULT_MODEL or {}
    prov = role_provider(role)
    fam = provider_family(prov)
    mid = (model or model_for(role) or "").strip()      # CEO's per-role model (or the default)
    shared = (dm.get("model") or "").strip()            # operator's shared/override model

    if prov == "openai_compatible":
        m = shared or (mid if MODEL_REGISTRY.get(mid, {}).get("provider") == "openai" else "gpt-4o")
        base = (dm.get("base_url") or "").strip() or _OPENAI_CLOUD_BASE
        if has_images and not _openai_is_vision(m):
            return _vision_fallback()
        return (m, "openai", base)

    if fam == "openai":
        # CEO-assigned OpenAI model wins; else the operator's shared model; else a sensible default.
        if MODEL_REGISTRY.get(mid, {}).get("provider") == "openai":
            m = mid
        else:
            m = shared or "gpt-4o"
        if has_images and not _openai_is_vision(m):
            return _vision_fallback()
        return (m, "openai", _OPENAI_CLOUD_BASE)

    # anthropic family (Claude — subscription or API): use the assigned Claude model, else Opus.
    cm = mid if MODEL_REGISTRY.get(mid, {}).get("provider") == "anthropic" else CEO_MODEL
    if has_images and not MODEL_REGISTRY.get(cm, {}).get("vision"):
        return _vision_fallback()
    return (cm, "anthropic", "")


# --------------------------------------------------------------------------
# Provider dispatch — the seam for adding/swapping an LLM backend. Each provider
# is a callable with a uniform signature; call_agent looks one up by name, so a
# new backend is added by registering it here — call_agent's routing never changes.
# --------------------------------------------------------------------------
_PROVIDERS = {}


def register_provider(name, fn):
    """Register an LLM backend. ``fn(model_id, system, user_prompt, messages, images,
    base_url, json_mode, temperature, max_tokens) -> str``. Replaces any existing entry."""
    _PROVIDERS[name] = fn


def _provider_anthropic(model_id, system, user_prompt, messages, images, base_url,
                        json_mode, temperature, max_tokens):
    return _call_anthropic(model_id, system, user_prompt, messages, images, max_tokens)


def _provider_openai(model_id, system, user_prompt, messages, images, base_url,
                     json_mode, temperature, max_tokens):
    return _call_openai_compatible(base_url, model_id, system, user_prompt, messages,
                                   images, json_mode, temperature, max_tokens)


register_provider("anthropic", _provider_anthropic)
register_provider("openai", _provider_openai)


def call_agent(role, system_prompt, user_prompt=None, messages=None, model=None,
               json_mode=False, images=None, max_tokens=16000, temperature=0.15):
    """
    Run `role` on its assigned model.

    Pass a single `user_prompt` (optionally with `images`), or a full `messages`
    list for multi-turn. `system_prompt` is prefixed with the role's persona.
    Returns the model's text. On any failure, returns an "...Error..." string so
    the UI can surface it without crashing.

    The provider/model is chosen by `_resolve_call`, which honours the global default
    model from Settings (applied to ALL agents) and falls back to a vision-capable
    model for image calls when the chosen model can't see.
    """
    model_id, provider, base_url = _resolve_call(role, model, images is not None)

    full_system = system_prompt or ""
    persona = persona_for(role)
    if persona:
        full_system = f"{persona}\n\n{full_system}".strip()

    if provider == "anthropic":
        # Route this role's plain text calls through the user's Claude subscription (Claude
        # Code) when that role is set to the subscription. Vision and multi-turn stay on the
        # API. Fall back to the API if Claude Code is unavailable / returns no JSON.
        on_subscription = role_provider(role) == "claude_subscription"
        if (on_subscription and images is None and messages is None):
            out = _call_claude_code(full_system, user_prompt, model=model_id)
            if out and (not json_mode or safe_json(out, default=_NO_JSON) is not _NO_JSON):
                return out
    # Dispatch to the registered provider backend (an unknown provider falls back to the
    # OpenAI-compatible client, matching the historical default).
    backend = _PROVIDERS.get(provider) or _PROVIDERS["openai"]
    return backend(model_id, full_system, user_prompt, messages, images, base_url,
                   json_mode, temperature, max_tokens)


# Engine error type + the call_agent failure-sentinel detector live in engine.errors.
from engine.errors import EngineError, is_engine_error, _ENGINE_ERROR_RE  # noqa: F401


def _call_anthropic(model_id, system, user_prompt, messages, images, max_tokens):
    client = anthropic_client()
    if client is None:
        return (f"Engine Error ({model_id}): Anthropic is not configured. "
                "Install the 'anthropic' package and set ANTHROPIC_API_KEY.")
    if messages is None:
        messages = [{"role": "user", "content": _anthropic_user_content(user_prompt, images)}]
    kwargs = dict(model=model_id, max_tokens=max_tokens, system=system,
                  thinking={"type": "adaptive"}, messages=messages)  # adaptive, no budget_tokens
    try:
        # Large outputs must stream to avoid the SDK's HTTP timeout. The default cap is 16000,
        # so stream at >= 16000 (not just > 16000) to cover the common full-size calls too.
        if max_tokens >= 16000:
            with client.messages.stream(**kwargs) as stream:
                resp = stream.get_final_message()
        else:
            resp = client.messages.create(**kwargs)
        text = "".join(b.text for b in resp.content if b.type == "text")
        _maybe_record_cost(model_id, getattr(resp, "usage", None))
        return text
    except Exception as e:
        return f"Claude Engine Error ({model_id}): {e}"


# --------------------------------------------------------------------------
# Optional: run the cloud (text) agents on the operator's Claude subscription
# via Claude Code (`claude -p`) instead of the metered API. Opt-in; falls back
# to the API on any failure. Vision/image calls always use the API.
# --------------------------------------------------------------------------
_SUBSCRIPTION_ROLES = set()   # roles whose cloud text calls run on the Claude Code subscription
_NO_JSON = object()


def set_subscription_roles(roles):
    """Set which agent roles route their (text) Anthropic calls through the
    Claude Code subscription instead of the metered API."""
    global _SUBSCRIPTION_ROLES
    _SUBSCRIPTION_ROLES = set((r or "").upper() for r in (roles or []))


def subscription_roles():
    """Routable roles currently set to run on the Claude subscription (from per-agent Settings)."""
    return [r for r in ROUTABLE_ROLES if role_provider(r) == "claude_subscription"]


def _claude_launcher():
    """Full path to the Claude Code CLI. `shutil.which` honours PATHEXT, so on
    Windows it resolves the `claude.CMD` shim — which `subprocess` can then launch
    by full path without a shell. Bare `"claude"` in an arg list fails with
    WinError 2 because CreateProcess doesn't apply PATHEXT. None if not installed."""
    return shutil.which("claude")


def _cc_model_alias(model_id):
    """Map an Anthropic model id to the alias Claude Code's --model accepts (opus/sonnet/
    haiku). Returns '' for unknown ids so we let Claude Code pick its default."""
    m = (model_id or "").lower()
    if "opus" in m:
        return "opus"
    if "sonnet" in m:
        return "sonnet"
    if "haiku" in m:
        return "haiku"
    return ""


def _call_claude_code(system, user_prompt, timeout=300, model=None):
    """Answer a single text prompt via Claude Code on the user's subscription.
    Runs headless in an empty temp dir (so it can't touch real files) with the
    API key stripped from the env (forces subscription auth). `model` (if it maps to
    a known alias) selects the subscription model so the Fast/Quality toggle takes
    effect on the subscription. Returns text or None on any failure (caller falls back)."""
    exe = _claude_launcher()
    if not exe:
        return None
    # Claude Code is an agent: without this it may CREATE files and print a summary
    # ("Done. I built…") instead of returning the content. Force a single-shot text reply.
    directive = ("You are a single-shot generator answering one request. Respond with ONLY the requested "
                 "output as your message text — no preamble, no explanation, no commentary. Do NOT use any "
                 "tools and do NOT create, write, or edit any files.")
    sys2 = f"{directive}\n\n{system}".strip() if system else directive
    prompt = f"{sys2}\n\n{user_prompt}"
    env = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}
    workdir = tempfile.mkdtemp(prefix="osw_cc_")
    cmd = [exe, "-p", "--output-format", "text", "--dangerously-skip-permissions"]
    alias = _cc_model_alias(model)
    if alias:
        cmd += ["--model", alias]
    try:
        # Prompt via stdin (not argv): avoids the ~32K Windows command-line limit
        # for large planning prompts and any shell-quoting issues.
        r = subprocess.run(
            cmd, input=prompt, cwd=workdir, env=env, capture_output=True, text=True,
            encoding="utf-8", errors="replace", timeout=timeout)
        out = _demojibake((r.stdout or "").strip())
        return out if (r.returncode == 0 and out) else None
    except Exception:
        return None
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def _call_openai_compatible(base_url, model_id, system, user_prompt, messages,
                            images, json_mode, temperature, max_tokens=16000):
    client = openai_client(base_url)
    if messages is None:
        messages = [{"role": "user", "content": _openai_user_content(user_prompt, images)}]
    chat = [{"role": "system", "content": system}] + messages
    kwargs = {"model": model_id, "messages": chat, "temperature": temperature,
              "max_tokens": max_tokens}
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    try:
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message.content
    except Exception as e:
        return f"Local Engine Error ({model_id} @ {base_url}): {e}"


def safe_json(raw, default=None):
    """Best-effort JSON parse that tolerates code fences and stray prose."""
    if default is None:
        default = {}
    if not raw:
        return default
    text = raw.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
    try:
        return json.loads(text)
    except Exception:
        for open_c, close_c in (("{", "}"), ("[", "]")):
            start = text.find(open_c)
            end = text.rfind(close_c)
            if start != -1 and end != -1 and end > start:
                try:
                    return json.loads(text[start:end + 1])
                except Exception:
                    continue
    return default


def encode_image(file_name, data_bytes):
    """Turn raw image bytes into the {media_type, b64} dict call_agent expects."""
    ext = os.path.splitext(file_name)[1].lower().lstrip(".")
    media = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
             "gif": "image/gif", "webp": "image/webp"}.get(ext, "image/png")
    return {"media_type": media, "b64": base64.standard_b64encode(data_bytes).decode("utf-8")}


_DOC_TEXT_EXTS = (".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log")


def extract_document_text(filename, data_bytes):
    """Extract plain text from an uploaded brief (PDF/Word/HTML/text). Returns
    (text, note); note is '' on success or a short human reason. Never raises."""
    ext = os.path.splitext(filename or "")[1].lower()
    try:
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except Exception:
                return "", "PDF support not installed (pip install pypdf)"
            reader = PdfReader(BytesIO(data_bytes))
            text = "\n".join((pg.extract_text() or "") for pg in reader.pages).strip()
            return (text, "") if text else ("", "scanned/empty PDF — paste the text instead")
        if ext == ".docx":
            try:
                import docx
            except Exception:
                return "", "Word support not installed (pip install python-docx)"
            doc = docx.Document(BytesIO(data_bytes))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(x for x in parts if x).strip()
            return (text, "") if text else ("", "empty document")
        if ext in (".html", ".htm"):
            html = data_bytes.decode("utf-8", errors="ignore")
            html = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\1>", " ", html)
            return re.sub(r"\s+", " ", re.sub(r"(?s)<[^>]+>", " ", html)).strip(), ""
        if ext in _DOC_TEXT_EXTS or ext == "":
            return data_bytes.decode("utf-8", errors="ignore").strip(), ""
        return "", f"unsupported file type ({ext})"
    except Exception as e:
        return "", f"could not read {os.path.basename(filename or '')}: {e}"


def _demojibake(s):
    """Repair UTF-8-decoded-as-Windows-1252 corruption (e.g. 'SÃ¶hne'→'Söhne',
    'â€“'→'—'). Only changes a string that contains the tell-tale bytes AND round-trips
    cleanly; otherwise returns it unchanged. Safe + idempotent (no-op on clean text)."""
    if not s or not isinstance(s, str):
        return s
    if "Ã" not in s and "â€" not in s and "Â" not in s:
        return s
    try:
        fixed = s.encode("cp1252", "strict").decode("utf-8", "strict")
        return fixed if fixed != s else s
    except (UnicodeEncodeError, UnicodeDecodeError):
        return s


# --------------------------------------------------------------------------
# Per-project local workspace
# --------------------------------------------------------------------------
def project_paths(base_dir, project_name):
    root = os.path.join(base_dir, project_name)
    return {
        "root": root,
        "src": os.path.join(root, "src"),
        "history": os.path.join(root, "history"),
        "changelog": os.path.join(root, "history", "changelog.md"),
        "bugs": os.path.join(root, "bugs.md"),
        "prd": os.path.join(root, "prd.md"),
        "plan": os.path.join(root, "plan.md"),
    }


def delete_project(base_dir, name):
    """Permanently delete a project's folder and ALL its files. Refuses if a build is still
    running for the project, or if the resolved folder would escape the workspace (so a bad
    name can never rmtree the workspace itself or an arbitrary path). Returns (ok, message).

    Uses ignore_errors so a Windows file lock (e.g. a still-open log handle) degrades to a
    partial delete reported honestly, rather than raising."""
    nm = (name or "").strip()
    if not nm or nm in (".", "..") or "/" in nm or "\\" in nm or os.path.isabs(nm):
        return False, "Invalid project name."
    if _any_build_active(nm):
        return False, "A build is still running for this project. Stop it first, then delete."
    base_real = os.path.realpath(base_dir)
    root = os.path.realpath(project_paths(base_dir, nm)["root"])
    if root == base_real or os.path.dirname(root) != base_real:
        return False, "Refusing to delete a folder outside the workspace."
    if not os.path.exists(root):
        return False, "Project folder not found (it may already be deleted)."
    shutil.rmtree(root, ignore_errors=True)
    with _JOBS_LOCK:                                  # drop any finished job handles for this name
        for k in (f"cc:{nm}", f"cct:{nm}", f"plan:{nm}", nm):
            _JOBS.pop(k, None)
    if os.path.exists(root):
        return False, (f"Deleted most of '{nm}', but some files were locked and remain. "
                       "Close anything open in that folder and try again.")
    return True, f"Deleted '{nm}' and all its files."


def init_project(base_dir, project_name):
    """Create the per-project folder structure and seed log files."""
    p = project_paths(base_dir, project_name)
    for key in ("root", "src", "history"):
        os.makedirs(p[key], exist_ok=True)
    if not os.path.exists(p["changelog"]):
        with open(p["changelog"], "w", encoding="utf-8") as f:
            f.write(f"# Change History — {project_name}\n\n"
                    "Append-only log of every decision, change and milestone.\n")
    if not os.path.exists(p["bugs"]):
        with open(p["bugs"], "w", encoding="utf-8") as f:
            f.write(f"# Bug Knowledge Base — {project_name}\n\n"
                    "QA logs every bug here. It is re-read on each QA pass so the "
                    "same bug is never shipped twice.\n")
    return p


def _now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def log_history(base_dir, project_name, agent, message):
    """Append a timestamped entry to the project's change history."""
    p = init_project(base_dir, project_name)
    with open(p["changelog"], "a", encoding="utf-8") as f:
        f.write(f"\n### {_now()} — {agent}\n{message}\n")
    return True


def log_bug(base_dir, project_name, title, detail, severity="medium",
            category="CODE", status="open"):
    """Append a structured bug entry to the project's bug knowledge base."""
    p = init_project(base_dir, project_name)
    with open(p["bugs"], "a", encoding="utf-8") as f:
        f.write(
            f"\n## [{category}] {title}\n"
            f"- **Logged:** {_now()}\n"
            f"- **Severity:** {severity}\n"
            f"- **Status:** {status}\n"
            f"- **Detail:** {detail}\n"
        )
    return True


def read_bugs(base_dir, project_name):
    """Return the bug knowledge base so QA can check for regressions."""
    p = project_paths(base_dir, project_name)
    if os.path.exists(p["bugs"]):
        with open(p["bugs"], "r", encoding="utf-8") as f:
            return f.read()
    return "No bugs logged yet."


def save_artifact(base_dir, project_name, file_name, content, subdir=None):
    """Overwrite a current-state artifact. subdir='src' drops it under src/."""
    p = init_project(base_dir, project_name)
    target_dir = p["src"] if subdir == "src" else p["root"]
    os.makedirs(target_dir, exist_ok=True)
    full_path = os.path.join(target_dir, file_name)
    with open(full_path, "w", encoding="utf-8") as f:
        f.write(content or "")
    return full_path


# --------------------------------------------------------------------------
# Project persistence (so all projects survive restarts and show on the dashboard)
# --------------------------------------------------------------------------
def save_project_state(base_dir, name, state):
    """Write resumable project state to project.json. MERGES into existing state so a
    partial update never wipes keys written elsewhere (e.g. state_summary, spec_score,
    prd/plan content)."""
    p = init_project(base_dir, name)
    existing = load_project_state(base_dir, name) or {}
    out = {**existing, **state}
    out["name"] = name
    out["created_at"] = existing.get("created_at", _now())
    out["updated_at"] = _now()
    with open(os.path.join(p["root"], "project.json"), "w", encoding="utf-8") as f:
        json.dump(out, f, indent=2)
    return out


def load_project_state(base_dir, name):
    path = os.path.join(base_dir, name, "project.json")
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None
    return None


def list_projects(base_dir):
    """All projects with a project.json, newest first."""
    out = []
    if not os.path.isdir(base_dir):
        return out
    for entry in os.listdir(base_dir):
        st = load_project_state(base_dir, entry)
        if st:
            out.append(st)
    out.sort(key=lambda s: s.get("updated_at", ""), reverse=True)
    return out


# --------------------------------------------------------------------------
# Version roadmap (Jira-like): features bucketed into v1 / v2 / v3 / backlog
# --------------------------------------------------------------------------
def _empty_roadmap():
    return {"versions": ["v1 (MVP)", "v2", "v3", "Backlog"], "items": [], "next_id": 0}


def roadmap_path(base_dir, name):
    return os.path.join(base_dir, name, "roadmap.json")


def _read_json(path, default):
    """Read a JSON file, returning `default` when it's missing or unreadable. (One audited reader for
    the simple artifact readers, instead of repeating the same open/json.load/except block.)"""
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return default


def read_roadmap(base_dir, name):
    return _read_json(roadmap_path(base_dir, name), _empty_roadmap())


def roadmap_to_md(rm):
    lines = ["# Version Roadmap", ""]
    for v in rm.get("versions", []):
        lines.append(f"## {v}")
        vitems = [it for it in rm.get("items", []) if it.get("version") == v]
        if not vitems:
            lines.append("_(no items)_")
        for it in vitems:
            pr = it.get("priority", "")
            detail = f" — {it.get('detail', '')}" if it.get("detail") else ""
            lines.append(f"- **{it.get('title', '')}**" + (f" _({pr})_" if pr else "") + detail)
        lines.append("")
    return "\n".join(lines)


def write_roadmap(base_dir, name, rm):
    init_project(base_dir, name)
    with open(roadmap_path(base_dir, name), "w", encoding="utf-8") as f:
        json.dump(rm, f, indent=2)
    save_artifact(base_dir, name, "roadmap.md", roadmap_to_md(rm))   # human-readable mirror
    return rm


# --------------------------------------------------------------------------
# Blueprint: the file + function hierarchy (the "whiteboard" structure)
# --------------------------------------------------------------------------
def _empty_blueprint():
    return {"summary": "", "files": []}


def blueprint_path(base_dir, name):
    return os.path.join(base_dir, name, "blueprint.json")


def read_blueprint(base_dir, name):
    return _read_json(blueprint_path(base_dir, name), _empty_blueprint())


def write_blueprint(base_dir, name, bp):
    init_project(base_dir, name)
    with open(blueprint_path(base_dir, name), "w", encoding="utf-8") as f:
        json.dump(bp, f, indent=2)
    return bp


def sync_from_folder(base_dir, name, model=None):
    """Re-read the actual project folder and refresh the platform's view:
    regenerate the file/function blueprint from the real code and write a
    current-state summary (used to inform v2/v3 planning). Returns the summary."""
    set_cost_context(base_dir, name, "sync")   # attribute this analysis's API spend to the ledger
    files = _list_code_files(base_dir, name)
    code = _read_project_code(base_dir, name, files, cap=150000) if files else ""
    if not code.strip():
        return "No code files found in the project folder to sync."
    out = safe_json(call_agent("CTO", "Respond with valid JSON only.", (
        "Read the ACTUAL project code below and produce: (1) a concise CURRENT-STATE summary of "
        "what is implemented (key modules, what works, notable gaps), and (2) a file & function "
        "blueprint. Return JSON only: {\"summary\": str, \"files\": [{\"path\": str, \"purpose\": "
        "str, \"functions\": [str]}]}\n\n"
        f"=== PROJECT FILES ===\n{code}"),
        model=(model or model_for("CTO", {})), json_mode=True, max_tokens=10000),
        default=_empty_blueprint())
    write_blueprint(base_dir, name, out)
    st = load_project_state(base_dir, name) or {}
    st["state_summary"] = out.get("summary", "")
    st["synced_at"] = _now()
    st["file_count"] = len(files)
    save_project_state(base_dir, name, st)
    if not _any_build_active(name):          # reconcile the dashboard chip, but never stomp a LIVE build
        # A sync means the operator re-read working code from the folder — clear any stale build
        # outcome ("Build failed"/interrupted/incomplete) so the card shows a neutral "Synced", not a
        # past dashboard build's result.
        write_status(base_dir, name, running=False, error=None, interrupted=False,
                     incomplete=False, finished_at=None, synced=True, action="Synced from folder")
    log_history(base_dir, name, "CTO",
                f"Synced from folder: {len(files)} file(s); blueprint + state summary refreshed.")
    return out.get("summary", "")


# --------------------------------------------------------------------------
# Reference material the operator attaches (mockups, API docs, sample data…)
# --------------------------------------------------------------------------
_REF_TEXT_EXTS = (".md", ".txt", ".json", ".csv", ".yaml", ".yml", ".xml", ".html",
                  ".ts", ".js", ".py", ".sql")


def references_dir(base_dir, name):
    d = os.path.join(project_paths(base_dir, name)["root"], "references")
    os.makedirs(d, exist_ok=True)
    return d


def list_references(base_dir, name):
    d = references_dir(base_dir, name)
    return sorted(os.listdir(d)) if os.path.isdir(d) else []


def save_reference(base_dir, name, filename, data_bytes):
    safe = os.path.basename(filename)
    with open(os.path.join(references_dir(base_dir, name), safe), "wb") as f:
        f.write(data_bytes)
    return safe


def read_text_references(base_dir, name, cap=20000):
    """Concatenate text-based reference files (capped) to ground planning."""
    d = references_dir(base_dir, name)
    if not os.path.isdir(d):
        return ""
    chunks, total = [], 0
    for fn in sorted(os.listdir(d)):
        if not fn.lower().endswith(_REF_TEXT_EXTS):
            continue
        try:
            body = open(os.path.join(d, fn), encoding="utf-8", errors="ignore").read()
        except Exception:
            continue
        chunk = f"--- REFERENCE: {fn} ---\n{body}\n"
        if total + len(chunk) > cap:
            chunks.append(f"--- REFERENCE: {fn} (truncated) ---\n{body[:max(0, cap - total)]}\n")
            break
        chunks.append(chunk)
        total += len(chunk)
    return "\n".join(chunks)


# --------------------------------------------------------------------------
# Data model & API contract (first-class, editable)
# --------------------------------------------------------------------------
def _empty_datamodel():
    return {"entities": [], "endpoints": []}


def datamodel_path(base_dir, name):
    return os.path.join(base_dir, name, "datamodel.json")


def read_datamodel(base_dir, name):
    return _read_json(datamodel_path(base_dir, name), _empty_datamodel())


def datamodel_to_md(dm):
    lines = ["# Data Model & API Contract", "", "## Entities"]
    for e in dm.get("entities", []):
        lines.append(f"### {e.get('name', '')}")
        for f in e.get("fields", []):
            note = f" — {f.get('notes', '')}" if f.get("notes") else ""
            lines.append(f"- `{f.get('name', '')}`: {f.get('type', '')}{note}")
        lines.append("")
    lines.append("## API Endpoints")
    for ep in dm.get("endpoints", []):
        note = f" — {ep.get('notes', '')}" if ep.get("notes") else ""
        lines.append(f"- **{ep.get('method', '')} {ep.get('path', '')}**{note}")
        if ep.get("request"):
            lines.append(f"  - request: {ep['request']}")
        if ep.get("response"):
            lines.append(f"  - response: {ep['response']}")
    return "\n".join(lines)


def write_datamodel(base_dir, name, dm):
    init_project(base_dir, name)
    with open(datamodel_path(base_dir, name), "w", encoding="utf-8") as f:
        json.dump(dm, f, indent=2)
    save_artifact(base_dir, name, "datamodel.md", datamodel_to_md(dm))   # mirror for the AI
    return dm


# --------------------------------------------------------------------------
# Live web-research grounding (Anthropic server-side web_search tool)
# --------------------------------------------------------------------------
def research_topic(query, max_tokens=4000):
    """Grounded research via Claude's web_search tool. Returns text (with sources)
    or None if the API/tool is unavailable (planning then proceeds without it)."""
    client = anthropic_client()
    if client is None:
        return None
    try:
        resp = client.messages.create(
            model=CEO_MODEL, max_tokens=max_tokens,
            tools=[{"type": "web_search_20260209", "name": "web_search"}],
            messages=[{"role": "user", "content": query}])
        text = "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")
        return text.strip() or None
    except Exception:
        return None


# --------------------------------------------------------------------------
# Lessons-learned feedback loop (per project + cross-project memory)
# --------------------------------------------------------------------------
def lessons_global_path(base_dir):
    return os.path.join(base_dir, "_lessons_global.md")


def read_global_lessons(base_dir, cap=6000):
    p = lessons_global_path(base_dir)
    if os.path.exists(p):
        try:
            return open(p, encoding="utf-8").read()[-cap:]
        except Exception:
            return ""
    return ""


def read_project_lessons(base_dir, name):
    """Project-specific lessons.md TEXT (the per-project post-mortem captured at sign-off)."""
    return _read_artifact_text(base_dir, name, "lessons.md")


def _read_artifact_text(base_dir, name, fn):
    p = os.path.join(base_dir, name, fn)
    if os.path.exists(p):
        try:
            return open(p, encoding="utf-8").read()
        except Exception:
            return ""
    return ""


def record_lessons(base_dir, name, model=None):
    """Summarize what QA/red-team/build surfaced into lessons.md and append to the
    cross-project memory, so future plans avoid the same mistakes."""
    set_cost_context(base_dir, name, "lessons")   # attribute this run's API spend to the ledger
    bugs = read_bugs(base_dir, name)
    redteam = _read_artifact_text(base_dir, name, "REDTEAM.md")
    st = read_status(base_dir, name) or {}
    summary = call_agent("PM", "", (
        "Extract concise LESSONS LEARNED from this project's QA bug log, red-team findings and build "
        "outcome: recurring bugs to avoid, what tripped up the build, and concrete guidance for next "
        "time. Output short markdown bullets.\n\n"
        f"=== QA BUGS ===\n{bugs[:6000]}\n\n=== RED-TEAM ===\n{redteam[:4000]}\n\n"
        f"=== PM/CTO REVIEW ===\n{st.get('pm_summary', '')}\n{st.get('cto_summary', '')}"),
        model=(model or CEO_MODEL), max_tokens=3000)
    save_artifact(base_dir, name, "lessons.md", summary)
    try:
        with open(lessons_global_path(base_dir), "a", encoding="utf-8") as f:
            f.write(f"\n\n## {name} ({_now()})\n{summary}\n")
    except Exception:
        pass
    log_history(base_dir, name, "PM", "Captured lessons learned (project + cross-project memory).")
    return summary


# Client-facing report export (build_client_report, _md_to_html_doc) moved to engine.reports.


# --------------------------------------------------------------------------
# Per-task Claude Code build (one focused run per v1 feature)
# --------------------------------------------------------------------------
def start_claude_code_build_tasks(base_dir, name):
    with _JOBS_LOCK:
        if _cc_build_alive(name):        # mutually exclusive with the one-shot build (same folder)
            return False
        t = threading.Thread(target=_run_cc_tasks, args=(base_dir, name), daemon=True)
        _JOBS["cct:" + name] = t
        t.start()
        return True


def _run_cc_tasks(base_dir, name):
    p = init_project(base_dir, name)
    root = p["root"]
    try:
        if not os.path.exists(os.path.join(root, "CLAUDE.md")):
            prepare_vscode_build(base_dir, name)
        write_status(base_dir, name, phase="build", engine="claude_code", running=True,
                     progress=5, action="Starting Claude Code (per-feature)…", started_at=_now(),
                     finished_at=None, error=None, cc_summary="")
        if not claude_code_available():
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         action="Claude Code not installed", finished_at=_now(),
                         error="Claude Code CLI not found. Install it with "
                               "`npm install -g @anthropic-ai/claude-code` and run `claude` once to log in.")
            return
        rm = read_roadmap(base_dir, name)
        v1 = (rm.get("versions") or ["v1 (MVP)"])[0]
        feats = [it for it in rm.get("items", []) if it.get("version") == v1] or \
                [{"title": "the v1/MVP scope", "detail": ""}]
        board = {"kind": "build", "phase": "Building per feature with Claude Code", "round": 0,
                 "tasks": [{"id": f"f{i}", "title": ft.get("title") or f"Feature {i + 1}", "state": "backlog",
                            "assignee": "Developer", "note": ""} for i, ft in enumerate(feats)]}
        write_board(base_dir, name, board)
        child_env = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}
        cc_exe = _claude_launcher()
        if not cc_exe:
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         action="Claude Code not found", finished_at=_now(),
                         error="Claude Code launcher not resolved (shutil.which('claude') failed).")
            return
        n = len(feats)
        for i, ft in enumerate(feats):
            if _build_cancelled(base_dir, name):    # operator stopped the build mid-way
                break
            _title = ft.get("title") or f"Feature {i + 1}"
            board["tasks"][i].update(state="dev", note="Claude Code implementing…")
            write_board(base_dir, name, board)
            write_status(base_dir, name, phase="build", engine="claude_code",
                         progress=int(5 + 90 * i / max(n, 1)),
                         action=f"Claude Code: {_title} ({i + 1}/{n})")
            prompt = (
                f"Implement ONLY this v1 feature into the existing project: '{_title}'"
                + (f" — {ft.get('detail', '')}" if ft.get("detail") else "") + ".\n"
                "Read CLAUDE.md, TECH_BRIEF.md, TECH_SPEC.md and ACCEPTANCE.md. Integrate with the "
                "existing code (don't break other features), satisfy this feature's acceptance "
                "criteria, install deps, run the app and fix errors before finishing.")
            log_path = os.path.join(p["history"], f"cc_task_{i}.log")
            try:
                rc = _run_cc(cc_exe, prompt, root, child_env, log_path, name, 1800,
                             _build_model_alias(base_dir, name))
            except subprocess.TimeoutExpired:
                board["tasks"][i].update(state="error", note="Timed out")   # a timeout is a failure
                write_board(base_dir, name, board)
                continue
            if rc == 0:                              # don't report success on a failed build
                board["tasks"][i].update(state="done", note="Built")
                log_history(base_dir, name, "DEVELOPER", f"Claude Code built feature: {_title}")
            else:
                board["tasks"][i].update(state="error", note=f"Build failed (exit {rc})")
                log_history(base_dir, name, "DEVELOPER",
                            f"Claude Code FAILED on feature: {_title} (exit {rc})")
            write_board(base_dir, name, board)
        if _build_cancelled(base_dir, name):        # don't overwrite the cancelled status with a tally
            return
        if not os.path.exists(os.path.join(root, "RUN_GUIDE.md")):
            try:
                generate_run_guide(base_dir, name)
            except Exception:
                pass
        built = sum(1 for t in board["tasks"] if t.get("state") == "done")
        failed = n - built
        if built == 0:                      # nothing built → hard error (gate sign-off), like the one-shot
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         progress=100, finished_at=_now(),
                         action="Build failed — nothing built",
                         error="No features built — every feature failed or timed out. Check the "
                               "per-feature logs, then retry the build.",
                         cc_summary=f"0/{n} features built.")
        elif failed:                        # partial — never report 'complete' when some failed/timed out
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         progress=100, finished_at=_now(), incomplete=True,
                         action=f"Built {built} of {n} features — {failed} failed. Check logs or retry.",
                         error=(f"{built}/{n} features built; {failed} failed or timed out — "
                                "see the per-feature logs, then retry the build."),
                         cc_summary=f"Built {built}/{n} feature(s); {failed} failed or timed out.")
        else:
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         progress=100, action="All v1 features built", finished_at=_now(),
                         cc_summary=f"Built {n} feature(s) task-by-task.")
    except Exception as e:
        write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                     action="Build failed", error=str(e))


# Brainstorm whiteboard DOT generation (build_graph_dot) moved to engine.graph.


# --------------------------------------------------------------------------
# Background development jobs (thread + status file + VS Code launch)
# --------------------------------------------------------------------------
_JOBS = {}
_JOBS_LOCK = threading.Lock()       # check-and-start must be atomic — never two builds per project
_CC_PROCS = {}                      # name -> the live Claude Code child Popen, so a build can be cancelled


def _build_model_alias(base_dir, name):
    """Claude Code --model alias (opus/sonnet/haiku) for the CEO-assigned DEVELOPER model — the
    autonomous build IS the Developer. Returns '' to let Claude Code use its subscription default."""
    mm = (load_project_state(base_dir, name) or {}).get("model_map") or {}
    return _cc_model_alias(model_for("DEVELOPER", mm))


def _run_cc(cc_exe, prompt, cwd, env, log_path, name, timeout, model_alias=""):
    """Run Claude Code as a CANCELLABLE child: register the Popen so cancel_build(name) can terminate it
    (which unblocks this worker thread so a fresh build can start). `model_alias` (opus/sonnet/haiku)
    runs the build on the CEO-assigned Developer model; '' uses Claude Code's default. Returns the exit
    code; raises subprocess.TimeoutExpired on timeout. Always deregisters the handle."""
    cmd = [cc_exe, "-p", "--dangerously-skip-permissions"]
    if model_alias:
        cmd += ["--model", model_alias]
    with open(log_path, "w", encoding="utf-8") as lf:
        proc = subprocess.Popen(cmd,
                                stdin=subprocess.PIPE, stdout=lf, stderr=subprocess.STDOUT,
                                text=True, encoding="utf-8", errors="replace", cwd=cwd, env=env)
        with _JOBS_LOCK:
            _CC_PROCS[name] = proc
        try:
            proc.communicate(input=prompt, timeout=timeout)
            return proc.returncode
        except subprocess.TimeoutExpired:
            try:
                proc.kill()
                proc.communicate()
            except Exception:
                pass
            raise
        finally:
            with _JOBS_LOCK:
                if _CC_PROCS.get(name) is proc:
                    _CC_PROCS.pop(name, None)


def _build_cancelled(base_dir, name):
    return bool((read_status(base_dir, name) or {}).get("cancelled"))


def cancel_build(base_dir, name):
    """Operator stop of a running SaaS build: mark the status cancelled and terminate the child Claude
    Code process so the worker thread unblocks (a fresh build can then start). The worker's finalize
    checks `cancelled` and won't overwrite this status with a build tally."""
    write_status(base_dir, name, phase="build", running=False, cancelled=True, progress=100,
                 action="Build cancelled", finished_at=_now(),
                 error="Build cancelled. Files created so far are kept — start a fresh build when you're ready.")
    with _JOBS_LOCK:
        proc = _CC_PROCS.get(name)
    if proc is not None:
        try:
            proc.terminate()
        except Exception:
            pass


def _cc_build_alive(name):
    """True if a Claude Code BUILD (one-shot 'cc:' OR per-feature 'cct:') is running for `name`.
    The two build buttons target the same project folder, so they must be mutually exclusive."""
    for k in ("cc:" + name, "cct:" + name):
        j = _JOBS.get(k)
        if j and j.is_alive():
            return True
    return False


_BUILD_ACTIVE_CHECKS = []   # extra (name)->bool liveness checks optionally registered by other modules


def register_build_active_check(fn):
    """Register a callable(name)->bool reporting whether a build worker for `name` is alive in THIS
    process, so orphan-reconciliation can cover builds owned by other modules without a circular
    import. (Unused in the app-only build; kept as a generic extension point.)"""
    if fn not in _BUILD_ACTIVE_CHECKS:
        _BUILD_ACTIVE_CHECKS.append(fn)

# --------------------------------------------------------------------------
# Per-project API cost ledger. Token usage is attributed to whichever project
# and phase is "active" on the current thread (set via set_cost_context).
# --------------------------------------------------------------------------
_cost_ctx = threading.local()
_cost_lock = threading.Lock()


def set_cost_context(base_dir, name, phase):
    _cost_ctx.base_dir, _cost_ctx.name, _cost_ctx.phase = base_dir, name, phase


def clear_cost_context():
    for a in ("base_dir", "name", "phase"):
        if hasattr(_cost_ctx, a):
            delattr(_cost_ctx, a)


def costs_path(base_dir, name):
    return os.path.join(base_dir, name, "costs.json")


def _empty_costs():
    return {"total_usd": 0.0, "input_tokens": 0, "output_tokens": 0, "calls": 0, "by_phase": {}}


def read_costs(base_dir, name):
    path = costs_path(base_dir, name)
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return _empty_costs()
    return _empty_costs()


def _model_cost(model_id, usage):
    """USD cost + (input+cache tokens, output tokens) for one Anthropic call."""
    inp = getattr(usage, "input_tokens", 0) or 0
    out = getattr(usage, "output_tokens", 0) or 0
    cr = getattr(usage, "cache_read_input_tokens", 0) or 0
    cw = getattr(usage, "cache_creation_input_tokens", 0) or 0
    in_rate, out_rate = PRICING.get(model_id, (0.0, 0.0))
    usd = (inp * in_rate + out * out_rate + cr * in_rate * 0.1 + cw * in_rate * 1.25) / 1_000_000
    return usd, inp + cr + cw, out


def _record_cost(base_dir, name, phase, usd, inp, out):
    with _cost_lock:
        c = read_costs(base_dir, name)
        c["total_usd"] = round(c.get("total_usd", 0) + usd, 6)
        c["input_tokens"] = c.get("input_tokens", 0) + inp
        c["output_tokens"] = c.get("output_tokens", 0) + out
        c["calls"] = c.get("calls", 0) + 1
        ph = c.setdefault("by_phase", {}).setdefault(
            phase, {"usd": 0.0, "input": 0, "output": 0, "calls": 0})
        ph["usd"] = round(ph["usd"] + usd, 6)
        ph["input"] += inp
        ph["output"] += out
        ph["calls"] += 1
        try:
            init_project(base_dir, name)
            with open(costs_path(base_dir, name), "w", encoding="utf-8") as f:
                json.dump(c, f, indent=2)
        except Exception:
            pass


def _maybe_record_cost(model_id, usage):
    """Record cost if a cost context is set on this thread (no-op otherwise)."""
    if usage is None:
        return
    name = getattr(_cost_ctx, "name", None)
    if not name:
        return
    usd, inp, out = _model_cost(model_id, usage)
    _record_cost(getattr(_cost_ctx, "base_dir", ""), name,
                 getattr(_cost_ctx, "phase", "other"), usd, inp, out)


# (The self-learning Playbook was removed.)


# --------------------------------------------------------------------------
# Global operator SETTINGS — API keys + the default agent model. Stored GLOBALLY
# (one settings.json beside the code), plaintext on disk (same trust model as .env).
# Applied into the running engine via apply_settings: keys
# go to the environment, the chosen model becomes the global default, and cached
# clients are dropped so changes take effect without a restart.
# --------------------------------------------------------------------------
_SETTINGS_PATH = os.path.join(_REPO_ROOT, "settings.json")
_settings_lock = threading.Lock()

_DEFAULT_SETTINGS = {
    "keys": {"anthropic": "", "openai": ""},
    "default_model": {"provider": "claude_subscription", "model": "", "base_url": ""},
    "agents": {},   # {ROLE: provider} per-agent overrides; empty = all use default_model.provider
}
# settings key -> environment variable the engine reads.
_SETTINGS_KEY_ENV = {
    "anthropic": "ANTHROPIC_API_KEY",
    "openai": "OPENAI_API_KEY",
}
# Providers offered in Settings. Keep claude_subscription first — it's the $0 default.
PROVIDER_CHOICES = ["claude_subscription", "anthropic", "openai", "openai_compatible"]


def _clean_agents(raw):
    """Keep only known routable roles mapped to known providers."""
    out = {}
    for r, p in (raw or {}).items():
        ru = str(r).upper()
        if ru in ROUTABLE_ROLES and p in PROVIDER_CHOICES:
            out[ru] = p
    return out


def settings_path():
    return _SETTINGS_PATH


def read_settings():
    """Global settings, merged onto safe defaults. Missing/unreadable file → defaults."""
    d = _read_json(_SETTINGS_PATH, None)
    if not isinstance(d, dict):
        d = {}
    keys = dict(_DEFAULT_SETTINGS["keys"]); keys.update(d.get("keys") or {})
    dm = dict(_DEFAULT_SETTINGS["default_model"]); dm.update(d.get("default_model") or {})
    if dm.get("provider") not in PROVIDER_CHOICES:
        dm["provider"] = "claude_subscription"
    return {"keys": keys, "default_model": dm, "agents": _clean_agents(d.get("agents"))}


def write_settings(d):
    """Persist settings atomically and apply them immediately. Returns the cleaned dict."""
    keys = dict(_DEFAULT_SETTINGS["keys"]); keys.update((d or {}).get("keys") or {})
    dm = dict(_DEFAULT_SETTINGS["default_model"]); dm.update((d or {}).get("default_model") or {})
    if dm.get("provider") not in PROVIDER_CHOICES:
        dm["provider"] = "claude_subscription"
    clean = {"keys": {k: (v or "").strip() for k, v in keys.items()}, "default_model": dm,
             "agents": _clean_agents((d or {}).get("agents"))}
    with _settings_lock:
        _atomic_write_json(_SETTINGS_PATH, clean)
    apply_settings(clean)
    return clean


def apply_settings(d=None):
    """Push settings into the running engine: API keys -> environment (only non-empty
    values, so an unset field never clobbers a key already provided via .env), the default
    model + per-agent providers -> the globals the router consults, and drop cached clients.
    Called at import and on every save so changes take effect live. The Claude Code BUILD
    still strips ANTHROPIC_API_KEY from its child env, so builds keep using the subscription."""
    global _DEFAULT_MODEL, _AGENT_PROVIDERS
    d = d if isinstance(d, dict) else read_settings()
    keys = d.get("keys") or {}
    for k, env in _SETTINGS_KEY_ENV.items():
        v = (keys.get(k) or "").strip()
        if v:
            os.environ[env] = v
    dm = dict(_DEFAULT_SETTINGS["default_model"]); dm.update(d.get("default_model") or {})
    if dm.get("provider") not in PROVIDER_CHOICES:
        dm["provider"] = "claude_subscription"
    _DEFAULT_MODEL = dm
    _AGENT_PROVIDERS = _clean_agents(d.get("agents"))
    reset_clients()


def default_model_label():
    """Short human label for the active default model, for status captions."""
    dm = _DEFAULT_MODEL or {}
    prov = dm.get("provider", "claude_subscription")
    if prov == "claude_subscription":
        return "Claude subscription ($0)"
    model = (dm.get("model") or "").strip() or "(model not set)"
    names = {"anthropic": "Claude API", "openai": "OpenAI", "openai_compatible": "OpenAI-compatible"}
    return f"{names.get(prov, prov)} · {model}"


# Apply persisted settings at import so the engine honours the saved default model + keys
# even before the UI runs. Safe when no settings.json exists.
try:
    apply_settings()
except Exception:
    pass


# (The self-learning Playbook rule engine was removed.)


def status_path(base_dir, name):
    return os.path.join(base_dir, name, "status.json")


def read_status(base_dir, name):
    path = status_path(base_dir, name)
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None
    return None


def _atomic_write_json(path, obj):
    """Write JSON atomically (unique temp in the same dir + os.replace) so the 2-3s UI fragments
    never read a half-written status.json/board.json. Falls back to a plain write on failure."""
    d = os.path.dirname(path) or "."
    os.makedirs(d, exist_ok=True)
    fd, tmp = tempfile.mkstemp(dir=d, prefix=".tmp_", suffix=".json")
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            json.dump(obj, f, indent=2)
        os.replace(tmp, path)            # atomic on Windows + POSIX (same volume)
    except Exception:
        try:
            os.remove(tmp)               # don't leak a temp file on a failed replace
        except Exception:
            pass
        with open(path, "w", encoding="utf-8") as f:   # last-resort non-atomic write
            json.dump(obj, f, indent=2)


def write_status(base_dir, name, **fields):
    init_project(base_dir, name)
    cur = read_status(base_dir, name) or {}
    cur.update(fields)
    if fields.get("running") is True:
        cur["interrupted"] = False      # a freshly (re)started run is no longer an 'interrupted' one
        cur["incomplete"] = False       # …nor a previously 'incomplete' (partial / unverified) one
        cur["synced"] = False           # …nor a 'synced-from-folder' one (a real build supersedes it)
        cur["cancelled"] = False        # …nor a previously 'cancelled' one
    _atomic_write_json(status_path(base_dir, name), cur)
    return cur


def build_active(name):
    """True iff a SaaS build / development / planning worker for `name` is alive in THIS process."""
    for key in ("cc:" + name, "cct:" + name, name, "plan:" + name):
        job = _JOBS.get(key)
        if job and job.is_alive():
            return True
    return False


def _any_build_active(name):
    """build_active() OR any registered cross-module liveness check."""
    if build_active(name):
        return True
    for fn in _BUILD_ACTIVE_CHECKS:
        try:
            if fn(name):
                return True
        except Exception:
            pass
    return False


def reconcile_status(base_dir, name):
    """Heal an orphaned 'running' status: if status.json says a build is running but no worker
    thread is alive in this process (the app or computer restarted mid-build, taking the child
    Claude Code process with it), flip it to a recoverable, retryable state so the UI is never
    stuck on a perpetual progress bar. No-op for a genuinely live build. Returns the status dict."""
    s = read_status(base_dir, name)
    if s and s.get("running") and not _any_build_active(name):
        s = write_status(base_dir, name, running=False, interrupted=True,
                         finished_at=(s.get("finished_at") or _now()),
                         action="Build interrupted",
                         error="The build was interrupted before it finished (the app or computer "
                               "restarted while it was running). Start it again to continue.")
    return s


# --- Kanban board (tasks flow Backlog → In Development → In QA → Done) -------
def board_path(base_dir, name):
    return os.path.join(base_dir, name, "board.json")


def read_board(base_dir, name):
    path = board_path(base_dir, name)
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None
    return None


def write_board(base_dir, name, board):
    init_project(base_dir, name)
    _atomic_write_json(board_path(base_dir, name), board)


def _build_board_from_roadmap(base_dir, name, state="backlog", note=""):
    """A build Kanban board seeded from the operator-approved v1/MVP roadmap features. Shared so the
    one-shot 'Develop with Claude Code' build shows the same kind of board as the per-feature build.
    Returns {kind, phase, round, tasks:[{id,title,state,assignee,note,purpose}]}."""
    rm = read_roadmap(base_dir, name)
    v1 = (rm.get("versions") or ["v1 (MVP)"])[0]
    feats = [it for it in (rm.get("items", []) or []) if it.get("version") == v1] or \
            [{"title": "the v1 / MVP scope", "detail": ""}]
    return {"kind": "build", "phase": "Building v1 with Claude Code", "round": 0,
            "tasks": [{"id": f"f{i}", "title": ft.get("title", f"Feature {i + 1}"),
                       "state": state, "assignee": "Developer", "note": note,
                       "purpose": ft.get("detail", "")} for i, ft in enumerate(feats)]}


def parse_code_files(text):
    """Extract (relpath, body) pairs from Developer output. Each file is a fenced
    code block preceded by a line 'FILE: <relative/path>'."""
    files = []
    if not text:
        return files
    pattern = re.compile(r"FILE:\s*(?P<path>[^\n`]+?)\s*```[^\n]*\n(?P<body>.*?)```", re.DOTALL)
    for m in pattern.finditer(text):
        path = m.group("path").strip().strip("`").strip()
        if path:
            files.append((path, m.group("body")))
    return files


def write_code_files(base_dir, name, files):
    """Write parsed files as a real tree under the project root (so VS Code shows them)."""
    p = init_project(base_dir, name)
    written = []
    for relpath, body in files:
        rel = relpath.replace("\\", "/").lstrip("/")
        if ".." in rel.split("/"):
            continue
        full = os.path.join(p["root"], *rel.split("/"))
        os.makedirs(os.path.dirname(full) or p["root"], exist_ok=True)
        with open(full, "w", encoding="utf-8") as f:
            f.write(body)
        written.append(rel)
    return written


def open_in_vscode(path):
    """Open a folder in VS Code (shell=True handles code.cmd on Windows)."""
    try:
        subprocess.Popen(f'code "{os.path.abspath(path)}"', shell=True)
        return True
    except Exception:
        return False


def prepare_vscode_build(base_dir, name):
    """Prepare the project so the operator's OWN AI coding agent (Claude Code,
    Cursor, Copilot) implements it in VS Code on their subscription. Writes a
    CLAUDE.md build brief (auto-read by Claude Code) and a user-facing
    BUILD_IN_VSCODE.md with exact terminal steps. Returns the steps markdown."""
    p = init_project(base_dir, name)
    root = p["root"]
    has = lambda fn: os.path.exists(os.path.join(root, fn))
    specs = [f"`{fn}`" for fn in ("prd.md", "plan.md", "TECH_BRIEF.md") if has(fn)]

    claude_md = (
        f"# {name} — build instructions for the AI coding agent\n\n"
        "Implement this project end to end. The specification documents are in this folder — read "
        "them fully before coding:\n\n"
        "- `prd.md` — Product Requirements Document (what to build)\n"
        "- `plan.md` — delivery plan (milestones, scope)\n"
        "- `TECH_BRIEF.md` — the CTO's engineering guidelines (architecture, clean architecture / "
        "layering, code-quality standards). **Follow these strictly.**\n"
        "- `roadmap.md` — the version roadmap. **Implement the v1/MVP scope ONLY.** Do NOT build "
        "v2/v3/backlog items now (treat them as future work).\n"
        "- `ACCEPTANCE.md` — acceptance criteria & examples. The build is DONE when these pass; "
        "use them as your checklist.\n"
        "- `TECH_SPEC.md` — constraints, decisions, data model and API contract. Honour these.\n"
        "- `references/` — operator-provided material (mockups, API docs, sample data). Use it.\n\n"
        "## Your task\n"
        "1. Implement the COMPLETE, runnable **v1/MVP** project: every v1 feature/module and every "
        "front-end page/component in its own file, plus config, package manifest(s), and "
        "`.env.example` for any required API keys.\n"
        "2. Follow the architecture and standards in `TECH_BRIEF.md`.\n"
        "3. Make sure it actually runs; install/verify dependencies.\n"
        "4. Write a `README.md` with EXACT, copy-pasteable setup and run commands, including the "
        "URL/port to open the front end in a browser.\n"
        "5. Finish with a short summary of what you built and how to run it.\n"
    )
    with open(os.path.join(root, "CLAUDE.md"), "w", encoding="utf-8") as f:
        f.write(claude_md)

    steps = (
        f"# Build “{name}” in VS Code with your AI coding agent\n\n"
        "The platform has prepared this folder with the full specs "
        f"({', '.join(specs) or 'the planning docs'}) and a `CLAUDE.md` build brief. Now let your "
        "own AI coding agent implement it **using your subscription** — no per-token API cost from "
        "this tool.\n\n"
        "## Option A — Claude Code (Opus 4.8) · recommended\n"
        "1. Install once (if you haven't):\n"
        "   ```powershell\n   npm install -g @anthropic-ai/claude-code\n   ```\n"
        "2. Open this folder in VS Code (use **Open in VS Code**, or run):\n"
        f"   ```powershell\n   code \"{root}\"\n   ```\n"
        "3. Open the VS Code terminal (Ctrl+`) and, from this folder, start Claude Code:\n"
        "   ```powershell\n   claude\n   ```\n"
        "   Sign in with your Claude subscription when prompted (this uses your plan, not an API key).\n"
        "4. Give it this prompt (Claude Code auto-reads `CLAUDE.md`, so “Build this project.” also "
        "works):\n"
        "   > Read CLAUDE.md, prd.md, plan.md and TECH_BRIEF.md, then implement the complete, "
        "runnable project following the engineering guidelines. When done, write a README.md with "
        "exact run commands and the URL to open the front end.\n\n"
        "5. When it finishes, open the generated **`README.md`** and run those commands to launch "
        "the app and view the front end.\n\n"
        "## Option B — Cursor / GitHub Copilot\n"
        "1. Open this folder in Cursor (or VS Code with Copilot Chat).\n"
        "2. In chat, reference `CLAUDE.md`, `prd.md`, `plan.md`, `TECH_BRIEF.md`.\n"
        "3. Ask it to implement the full project per those docs, then write a `README.md` with run "
        "commands.\n\n"
        "Either way, the coding runs on **your** AI subscription inside your editor.\n"
    )
    with open(os.path.join(root, "BUILD_IN_VSCODE.md"), "w", encoding="utf-8") as f:
        f.write(steps)
    log_history(base_dir, name, "PM",
                "Prepared the project for an external AI coding agent (CLAUDE.md + VS Code steps).")
    return steps


def _list_code_files(base_dir, name):
    """Project files that are actual code/assets (excludes platform metadata & docs)."""
    root = project_paths(base_dir, name)["root"]
    skip = {"project.json", "status.json", "board.json", "costs.json", "prd.md", "plan.md",
            "TECH_BRIEF.md", "RUN_GUIDE.md", "bugs.md", "CLAUDE.md", "BUILD_IN_VSCODE.md",
            ".env", ".env.example"}
    out = []
    for r, _, fs in os.walk(root):
        if (os.sep + "history") in r:
            continue
        for f in fs:
            if f in skip:
                continue
            out.append(os.path.relpath(os.path.join(r, f), root).replace("\\", "/"))
    return out


# Build-output verification — a build that exits 0 but created/changed NO source files did nothing
# (asked a question, hit a guardrail, no-op). Bracket the build subprocess with two snapshots and diff:
# during that window the ONLY writer is Claude Code (our status/board writes happen after; logs go to
# history/), so any created/changed source file is real work and an empty diff is a silent no-op.
_BUILD_SNAP_SKIP = {"project.json", "status.json", "board.json", "costs.json", "wtask.json"}


def snapshot_build_files(root):
    """{abspath: (size, mtime)} of candidate SOURCE files under `root` — excludes history/, dot-dirs
    (.git/.claude/…), dotfiles, and the platform status/metadata files. Never raises."""
    snap = {}
    try:
        for r, ds, fs in os.walk(root):
            ds[:] = [d for d in ds if not d.startswith(".") and d != "history"]
            for f in fs:
                if f.startswith(".") or f in _BUILD_SNAP_SKIP:
                    continue
                fp = os.path.join(r, f)
                try:
                    s = os.stat(fp)
                    snap[fp] = (s.st_size, int(s.st_mtime))
                except OSError:
                    continue
    except Exception:
        pass
    return snap


def build_made_changes(before, after):
    """True if any candidate source file was CREATED or its size/mtime CHANGED between two
    snapshot_build_files() maps — i.e. the build actually produced/edited source."""
    return any(before.get(fp) != sig for fp, sig in (after or {}).items())


def generate_run_guide(base_dir, name, model=None):
    """Author a detailed RUN_GUIDE.md from the project's actual files, on demand."""
    # Inherit the caller's cost phase (e.g. "build" when generated during a build); only default to
    # "build" when no context is set, so a standalone post-build regen isn't mis-billed.
    if not getattr(_cost_ctx, "name", None):
        set_cost_context(base_dir, name, "build")
    files = _list_code_files(base_dir, name)
    code = _read_project_code(base_dir, name, files, cap=120000) if files else "(no code files found)"
    prompt = (
        "Write a DETAILED, numbered RUN GUIDE (markdown) for a non-expert operator to run this "
        "project locally and see the front end. Include: exact prerequisites and install commands "
        "(copy-pasteable), how to start the app, the EXACT URL/port to open in a browser, how to "
        "verify it works, and troubleshooting. Base it ONLY on the actual files below.\n\n"
        f"=== PROJECT FILES ===\n{code}\n\nOutput markdown only."
    )
    guide = call_agent("PM", "", prompt, model=model or CEO_MODEL, max_tokens=6000)
    save_artifact(base_dir, name, "RUN_GUIDE.md", guide)
    log_history(base_dir, name, "PM", "Generated run guide on request from the delivered files.")
    return guide


# --------------------------------------------------------------------------
# Claude Code as the Developer — autonomous build on the operator's subscription
# --------------------------------------------------------------------------
def claude_code_available():
    """True if the Claude Code CLI is installed and callable."""
    try:
        r = subprocess.run("claude --version", shell=True, capture_output=True,
                           text=True, timeout=20)
        return r.returncode == 0
    except Exception:
        return False


def vscode_available():
    """True if the VS Code CLI ('code') is installed and callable."""
    try:
        r = subprocess.run("code --version", shell=True, capture_output=True,
                           text=True, timeout=20)
        return r.returncode == 0
    except Exception:
        return False


def node_available():
    """True if Node.js is installed (it provides `npm`, which Claude Code needs)."""
    try:
        r = subprocess.run("node --version", shell=True, capture_output=True,
                           text=True, timeout=20)
        return r.returncode == 0
    except Exception:
        return False


def refresh_path_from_registry():
    """On Windows, merge the current machine + user PATH from the registry into this process's
    environment. A tool installed AFTER the app started (e.g. Claude Code via `npm i -g`) lands on
    the PATH for *new* terminals, but the running app keeps its old PATH and can't see it. Calling
    this before a requirements re-check lets the app detect it without a restart. No-op off Windows."""
    if os.name != "nt":
        return
    try:
        import winreg
        vals = []
        for root, sub in ((winreg.HKEY_LOCAL_MACHINE,
                           r"SYSTEM\CurrentControlSet\Control\Session Manager\Environment"),
                          (winreg.HKEY_CURRENT_USER, "Environment")):
            try:
                with winreg.OpenKey(root, sub) as key:
                    raw, _ = winreg.QueryValueEx(key, "Path")
                    if raw:
                        vals.append(os.path.expandvars(raw))
            except OSError:
                pass
        seen, out = set(), []
        for p in os.pathsep.join(vals + [os.environ.get("PATH", "")]).split(os.pathsep):
            p = p.strip()
            if p and p.lower() not in seen:
                seen.add(p.lower())
                out.append(p)
        if out:
            os.environ["PATH"] = os.pathsep.join(out)
    except Exception:
        pass


def start_claude_code_build(base_dir, name):
    """Run Claude Code headlessly to implement the project (build → run → fix) on
    the operator's subscription. Returns False if a build is already running."""
    with _JOBS_LOCK:
        if _cc_build_alive(name):        # mutually exclusive with the per-feature build (same folder)
            return False
        t = threading.Thread(target=_run_claude_code_build, args=(base_dir, name), daemon=True)
        _JOBS["cc:" + name] = t
        t.start()
        return True


def _run_claude_code_build(base_dir, name):
    p = init_project(base_dir, name)
    root = p["root"]
    try:
        if not os.path.exists(os.path.join(root, "CLAUDE.md")):
            prepare_vscode_build(base_dir, name)   # ensure the build brief exists
        write_status(base_dir, name, phase="build", engine="claude_code", running=True,
                     progress=8, action="Starting Claude Code…", started_at=_now(),
                     finished_at=None, error=None, cc_summary="")

        if not claude_code_available():
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         action="Claude Code not installed", finished_at=_now(),
                         error="Claude Code CLI not found. Install it with "
                               "`npm install -g @anthropic-ai/claude-code`, then run `claude` once "
                               "in a terminal to sign in with your subscription, and rebuild.")
            return

        log_history(base_dir, name, "DEVELOPER",
                    "Claude Code (subscription) is implementing the project autonomously.")
        write_status(base_dir, name, phase="build", engine="claude_code", progress=30,
                     action="Claude Code: building, running and fixing the project…")
        # Live Kanban: seed from the approved v1 features so the operator watches them build
        # (the one-shot build is opaque per-feature, so all cards sit In Development until it ends).
        cc_board = None
        try:
            cc_board = _build_board_from_roadmap(base_dir, name, state="dev",
                                                 note="Claude Code building…")
            write_board(base_dir, name, cc_board)
        except Exception:
            cc_board = None
        prompt = (
            "Read CLAUDE.md, prd.md, plan.md, TECH_BRIEF.md, roadmap.md, ACCEPTANCE.md, "
            "TECH_SPEC.md and anything in references/ in this folder, then implement the COMPLETE, "
            "runnable **v1/MVP** scope from roadmap.md (do NOT build v2/v3/backlog items) following "
            "the engineering guidelines and TECH_SPEC.md, so that ACCEPTANCE.md passes. Create all "
            "source files, install dependencies, RUN the app and FIX every error until it starts cleanly. "
            "Ensure the dependency manifest lists every package used and module config is "
            "consistent. Then write/refresh README.md with exact, copy-pasteable run commands and "
            "the URL to open the front end. End with a short summary of what you built and how to run it."
        )
        log_path = os.path.join(p["history"], "claude_code_build.log")
        # Use the SUBSCRIPTION, not our API key: strip ANTHROPIC_API_KEY from the child env so
        # Claude Code authenticates with the logged-in plan (an env key would bill the API instead).
        child_env = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}
        cc_exe = _claude_launcher()
        if not cc_exe:
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         action="Claude Code not found", finished_at=_now(),
                         error="Claude Code launcher not resolved (shutil.which('claude') failed).")
            return
        _before = snapshot_build_files(root)   # to verify the build actually produces/edits source
        try:
            rc = _run_cc(cc_exe, prompt, root, child_env, log_path, name, 3600,
                         _build_model_alias(base_dir, name))
        except subprocess.TimeoutExpired:
            try:
                if cc_board:
                    for _t in cc_board["tasks"]:
                        _t.update(state="error", note="Timed out")
                    write_board(base_dir, name, cc_board)
            except Exception:
                pass
            write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                         action="Build timed out", finished_at=_now(),
                         error="Claude Code build exceeded 60 minutes and was stopped.")
            return
        if _build_cancelled(base_dir, name):   # operator stopped it — don't clobber the cancelled status
            return

        summary = ""
        try:
            summary = open(log_path, encoding="utf-8").read()[-4000:]
        except Exception:
            pass
        # Verify real output BEFORE writing any post-build artifact (RUN_GUIDE) — otherwise the run
        # guide we generate would itself count as "produced" and mask a no-op build.
        produced = build_made_changes(_before, snapshot_build_files(root))
        if not os.path.exists(os.path.join(root, "RUN_GUIDE.md")):
            try:
                generate_run_guide(base_dir, name)
            except Exception:
                pass
        log_history(base_dir, name, "DEVELOPER", f"Claude Code build finished (exit code {rc}).")
        if rc != 0:
            err = f"Claude Code exited with code {rc}. See history/claude_code_build.log."
            action = "Build failed"
        elif not produced:
            err = ("Build finished but produced no source files — open history/claude_code_build.log "
                   "to see what happened, then retry.")
            action = "Build produced no source files"
            log_history(base_dir, name, "DEVELOPER",
                        "Build exited cleanly but created/changed no source files.")
        else:
            err = None
            action = "Claude Code build complete"
        try:                                       # complete (or fail) every card on the live board
            if cc_board:
                _ok = err is None
                _note = "Built ✓" if _ok else ("No output" if rc == 0 else f"Build failed (exit {rc})")
                for _t in cc_board["tasks"]:
                    _t.update(state=("done" if _ok else "error"), note=_note)
                write_board(base_dir, name, cc_board)
        except Exception:
            pass
        write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                     progress=100, action=action, finished_at=_now(),
                     cc_summary=summary, error=err)
    except Exception as e:
        write_status(base_dir, name, phase="build", engine="claude_code", running=False,
                     action="Build failed", error=str(e))


def _strip_fences(text):
    """Clean a single-file model response: drop a leading 'FILE:' line and any
    surrounding ``` code fences, so we write pure file contents to disk."""
    t = (text or "").strip()
    lines = t.split("\n")
    if lines and lines[0].strip().upper().startswith("FILE:"):
        lines = lines[1:]
    if lines and lines[0].lstrip().startswith("```"):
        lines = lines[1:]
        if lines and lines[-1].rstrip().endswith("```"):
            lines = lines[:-1]
    return "\n".join(lines).strip() + "\n"


def _read_project_code(base_dir, name, rel_paths, cap=60000):
    """Concatenate the project's files (path-labelled) for QA/PM review, capped."""
    p = project_paths(base_dir, name)
    chunks, total = [], 0
    for rel in rel_paths:
        full = os.path.join(p["root"], *rel.replace("\\", "/").split("/"))
        if not os.path.isfile(full):
            continue
        try:
            body = open(full, encoding="utf-8").read()
        except Exception:
            continue
        chunk = f"FILE: {rel}\n{body}\n"
        if total + len(chunk) > cap:
            chunks.append(f"FILE: {rel}\n[omitted — review size cap reached]\n")
            break
        chunks.append(chunk)
        total += len(chunk)
    return "\n".join(chunks)


def write_env_files(base_dir, name, env_vars):
    """Create the project's .env.example (always) and .env (only if absent, so we
    never clobber keys the operator already filled) listing each external API key
    variable with a comment. Returns the list of variable names."""
    if not env_vars:
        return []
    p = init_project(base_dir, name)
    lines = ["# Environment variables for this project.",
             "# Fill in a value for each external service below, then save.", ""]
    names = []
    for v in env_vars:
        nm = (v.get("name") or "").strip()
        if not nm:
            continue
        names.append(nm)
        meta = " — ".join(x for x in [v.get("service", ""), v.get("description", "")] if x)
        if meta:
            lines.append(f"# {meta}")
        lines.append(f"{nm}=")
        lines.append("")
    if not names:
        return []
    content = "\n".join(lines)
    with open(os.path.join(p["root"], ".env.example"), "w", encoding="utf-8") as f:
        f.write(content)
    env_path = os.path.join(p["root"], ".env")
    if not os.path.exists(env_path):
        with open(env_path, "w", encoding="utf-8") as f:
            f.write(content)
    return names


# --------------------------------------------------------------------------
# Multi-agent, very detailed PRD + delivery plan
# --------------------------------------------------------------------------
def team_responsibilities_md(team, model_map):
    """Markdown header listing each agent's role, mandate and assigned model."""
    lines = ["## Team & Responsibilities", ""]
    for role in team:
        model = CEO_MODEL if role == "CEO" else model_for(role, model_map)
        lines.append(f"### {role}")
        lines.append(f"- **Model:** `{model}`")
        lines.append(f"- **Mandate:** {PERSONAS.get(role, '')}")
        lines.append("")
    return "\n".join(lines)


def build_plan_documents(base_dir, name, seed, discussion, team, model_map, on_progress=None,
                         options=None):
    """Each exec authors a deep section, then the CEO synthesizes exhaustive
    prd.md and plan.md. `options` toggles optional steps (research/blueprint/
    acceptance/tech_spec/datamodel/redteam). Returns (prd, plan)."""
    options = options or {}

    def opt(k, default=True):
        return options.get(k, default)

    def rep(pct, action):
        if on_progress:
            try:
                on_progress(pct, action)
            except Exception:
                pass

    refs = read_text_references(base_dir, name)
    if refs.strip():
        discussion = discussion + "\n\n=== REFERENCE MATERIAL (provided by the operator) ===\n" + refs

    _gl = read_global_lessons(base_dir)
    _pl = read_project_lessons(base_dir, name)
    if (_gl + _pl).strip():
        discussion += ("\n\n=== LESSONS LEARNED (apply these to avoid past mistakes) ===\n"
                       + _pl + "\n" + _gl)

    if opt("research", False):
        rep(4, "Researching market, APIs & best practices…")
        _res = research_topic(
            "Research to ground this software project's plan: current best-practice tech choices, "
            "up-to-date APIs/SDK versions, and notable competitors/approaches — with sources.\n\n"
            + seed[:3000])
        if _res:
            save_artifact(base_dir, name, "research.md", _res)
            discussion += "\n\n=== WEB RESEARCH (current, with sources) ===\n" + _res

    acc, techspec = "", ""
    team_md = team_responsibilities_md(team, model_map)
    # Planning Kanban — documents flow To-draft → Drafting → Done.
    board = {"kind": "planning", "phase": "Authoring the documents", "round": 0, "tasks": [
        {"id": "t0", "title": "Technical architecture", "state": "backlog", "assignee": "CTO", "note": ""},
        {"id": "t1", "title": "Product & market", "state": "backlog", "assignee": "CMO", "note": ""},
        {"id": "t2", "title": "Business case", "state": "backlog", "assignee": "CEO", "note": ""},
        {"id": "t3", "title": "PRD + Delivery plan", "state": "backlog", "assignee": "CEO", "note": "synthesis"},
        {"id": "t4", "title": "Engineering guidelines", "state": "backlog", "assignee": "CTO", "note": ""},
        {"id": "t5", "title": "Version roadmap (v1/v2/v3)", "state": "backlog", "assignee": "CEO", "note": ""},
        {"id": "t6", "title": "File & function blueprint", "state": "backlog", "assignee": "CTO", "note": ""},
        {"id": "t7", "title": "Acceptance criteria & examples", "state": "backlog", "assignee": "QA", "note": ""},
        {"id": "t8", "title": "Tech spec (constraints, data model, API)", "state": "backlog", "assignee": "CTO", "note": ""},
        {"id": "t9", "title": "Red-team review", "state": "backlog", "assignee": "SKEPTIC", "note": ""},
    ]}
    write_board(base_dir, name, board)
    _role_task = {"CTO": 0, "CMO": 1, "CEO": 2}
    steps = [
        ("CTO", 15, "CTO authoring the technical section…",
         "Write an exhaustive TECHNICAL section: system architecture, recommended tech stack with "
         "justification, data model/schema, APIs and third-party integrations, security and "
         "privacy, scalability and performance, infrastructure/devops/CI-CD, and a testing strategy."),
        ("CMO", 38, "CMO authoring the product & market section…",
         "Write an exhaustive PRODUCT & MARKET section: target market and competitors, detailed "
         "user personas, end-to-end user journeys, a prioritized feature list (MoSCoW), UX "
         "principles and key screens, and a go-to-market/positioning plan."),
        ("CEO", 60, "CEO authoring the business section…",
         "Write an exhaustive BUSINESS section: problem statement and business case, goals and "
         "non-goals, scope, measurable success metrics/KPIs, milestones and timeline, risks with "
         "mitigations, and a team/hiring plan."),
    ]
    drafts = {}
    # The three section drafts are independent, so run them concurrently to cut
    # wall-clock (each is a slow Opus call). The shared board.json is written once
    # before and once after — worker threads must NOT write it concurrently.
    for role, _pct, _action, _spec in steps:
        board["tasks"][_role_task[role]].update(state="dev", note="Drafting…")
    board["phase"] = "Drafting architecture, market & business sections (in parallel)…"
    write_board(base_dir, name, board)
    rep(20, "CTO, CMO & CEO drafting their sections in parallel…")

    # Worker threads don't inherit this thread's cost-attribution context — capture
    # and re-set it inside each worker so API costs are still booked to the project.
    _cc = (getattr(_cost_ctx, "base_dir", base_dir), getattr(_cost_ctx, "name", name),
           getattr(_cost_ctx, "phase", "planning"))

    def _draft(role, spec):
        set_cost_context(*_cc)
        prompt = (f"{spec}\nBe comprehensive, specific and descriptive — do not summarize.\n\n"
                  f"=== PROJECT BRIEF ===\n{seed}\n\n=== SCOPING DISCUSSION ===\n{discussion}\n\n"
                  "Output detailed markdown only, no preamble.")
        return call_agent(role, "", prompt,
                          model=(CEO_MODEL if role == "CEO" else model_for(role, model_map)),
                          max_tokens=12000)

    with ThreadPoolExecutor(max_workers=3) as _ex:
        _futs = {role: _ex.submit(_draft, role, spec) for role, _pct, _action, spec in steps}
        for role, fut in _futs.items():
            drafts[role] = fut.result()

    for role, _pct, _action, _spec in steps:
        board["tasks"][_role_task[role]].update(state="done", note="Drafted ✓")
    write_board(base_dir, name, board)
    rep(70, "Sections drafted — synthesizing the PRD & plan next…")

    board["tasks"][3].update(state="dev", note="Synthesizing…")
    board["phase"] = "Synthesizing the PRD & delivery plan"
    write_board(base_dir, name, board)
    rep(78, "CEO synthesizing the PRD & delivery plan…")
    synth = (
        "You are the CEO assembling the final project documents from the team's drafts and the "
        "scoping discussion. HONOUR the operator's Agree/Disagree decisions on suggestions. "
        "Produce TWO very detailed markdown documents separated by a single line containing "
        "exactly '===PLAN==='. First a Product Requirements Document, then a Delivery Plan "
        "(milestones, workstreams, task breakdown, dependencies, timeline, acceptance criteria). "
        "Be exhaustive and descriptive — do NOT shorten or summarize.\n\n"
        f"=== TEAM ===\n{team_md}\n\n=== CTO DRAFT ===\n{drafts['CTO']}\n\n"
        f"=== CMO DRAFT ===\n{drafts['CMO']}\n\n=== CEO DRAFT ===\n{drafts['CEO']}\n\n"
        f"=== DISCUSSION ===\n{discussion}\n\n"
        "Output the PRD markdown, then a line '===PLAN===', then the Plan markdown."
    )
    combined = call_agent("CEO", "", synth, model=CEO_MODEL, max_tokens=32000)
    if "===PLAN===" in combined:
        prd_body, plan_body = combined.split("===PLAN===", 1)
    else:
        prd_body, plan_body = combined, combined
    prd = team_md + "\n\n" + prd_body.strip()
    plan = team_md + "\n\n" + plan_body.strip()
    board["tasks"][3].update(state="done", note="Done ✓")
    write_board(base_dir, name, board)

    # ---- Post-synthesis stages, parallelized by dependency wave ----------
    # Worker threads don't inherit this thread's cost-attribution context, so each
    # re-sets it. Workers only write their own distinct artifact files; every
    # shared write (board.json, roadmap, changelog, project.json) stays on this
    # coordinating thread to avoid concurrent-write corruption.
    _cc2 = (getattr(_cost_ctx, "base_dir", base_dir), getattr(_cost_ctx, "name", name),
            getattr(_cost_ctx, "phase", "planning"))

    def _stage(fn):
        set_cost_context(*_cc2)
        return fn()

    # Wave A — engineering guidelines (t4) and the version roadmap (t5) depend
    # only on the PRD/plan/discussion, so run them concurrently.
    def _do_guidelines():
        return call_agent("CTO", "", (
            "You are the CTO. Write concise but concrete ENGINEERING GUIDELINES the Developer must "
            "follow while implementing this project. Cover: the recommended architecture and how to "
            "apply clean architecture / separation of concerns (layers and dependency direction); "
            "folder & file structure conventions; key design patterns; code-quality standards (clear "
            "naming, small focused functions, robust error handling, input validation at boundaries, "
            "no dead code); security practices; and the testing approach. Be specific and actionable "
            "— this is the rulebook applied to every file.\n\n"
            f"=== PLAN ===\n{plan[:6000]}\n\n=== PRD ===\n{prd[:8000]}\n\nOutput detailed markdown only."),
            model=model_for("CTO", model_map), max_tokens=6000)

    def _do_roadmap():
        return safe_json(call_agent("CEO", "Respond with valid JSON only.", (
            "From the scoping discussion and the PRD/plan, produce a VERSION ROADMAP. Decide the "
            "v1/MVP scope versus what should be deferred to v2, v3 or backlog — use the future / "
            "'phase 2' ideas the team and operator discussed. Return JSON only: "
            '{"versions": ["v1 (MVP)", "v2", "v3", "Backlog"], "items": [{"title": str, "detail": '
            'str, "version": str, "priority": "high|med|low"}]}\n\n'
            f"=== SCOPING DISCUSSION ===\n{discussion}\n\n=== PRD (excerpt) ===\n{prd[:6000]}"),
            model=CEO_MODEL, max_tokens=6000), default={})

    board["tasks"][4].update(state="dev", note="Writing guidelines…")
    board["tasks"][5].update(state="dev", note="Bucketing features by version…")
    board["phase"] = "Engineering guidelines & version roadmap (in parallel)…"
    write_board(base_dir, name, board)
    rep(82, "CTO writing guidelines while the CEO builds the roadmap…")
    with ThreadPoolExecutor(max_workers=2) as _ex:
        _fg = _ex.submit(_stage, _do_guidelines)
        _fr = _ex.submit(_stage, _do_roadmap)
        tech_brief = _fg.result()
        rmd = _fr.result()
    save_artifact(base_dir, name, "TECH_BRIEF.md", tech_brief)
    board["tasks"][4].update(state="done", note="Done ✓")

    rm_versions = rmd.get("versions") or ["v1 (MVP)", "v2", "v3", "Backlog"]
    rm_items = []
    for i, it in enumerate(rmd.get("items", []) or []):
        rm_items.append({"id": f"r{i}", "title": it.get("title", ""), "detail": it.get("detail", ""),
                         "version": it.get("version") or rm_versions[0],
                         "priority": it.get("priority", "")})
    write_roadmap(base_dir, name, {"versions": rm_versions, "items": rm_items,
                                   "next_id": len(rm_items)})
    log_history(base_dir, name, "CEO",
                f"Version roadmap created: {len(rm_items)} item(s) across {len(rm_versions)} versions.")
    board["tasks"][5].update(state="done", note="Done ✓")
    write_board(base_dir, name, board)

    _v1 = rm_versions[0] if rm_versions else "v1 (MVP)"
    _v1_feats = "\n".join(f"- {it['title']}: {it.get('detail', '')}"
                          for it in rm_items if it.get("version") == _v1)

    # Wave B — blueprint (t6), acceptance criteria (t7) and the tech spec + data
    # model (t8) are mutually independent now that the guidelines and v1 feature
    # list exist. Run whichever are enabled concurrently.
    def _do_blueprint():
        bp = safe_json(call_agent("CTO", "Respond with valid JSON only.", (
            "Produce a FILE & FUNCTION BLUEPRINT for the v1/MVP implementation — the files to create "
            "and the key functions in each, organized by folder. Return JSON only: {\"summary\": "
            "\"<one-paragraph architecture overview>\", \"files\": [{\"path\": \"relative/path\", "
            "\"purpose\": \"one line\", \"functions\": [\"fnName — what it does\"]}]}\n\n"
            f"=== PRD (excerpt) ===\n{prd[:7000]}\n\n"
            f"=== ENGINEERING GUIDELINES (excerpt) ===\n{tech_brief[:4000]}"),
            model=model_for("CTO", model_map), json_mode=True, max_tokens=8000), default=_empty_blueprint())
        write_blueprint(base_dir, name, bp)
        return ""

    def _do_acceptance():
        a = call_agent("QA", "", (
            "Write ACCEPTANCE CRITERIA for the v1/MVP features below. For EACH feature give testable "
            "Given/When/Then criteria and at least one concrete example (input → expected output). Make "
            "them unambiguous and verifiable. Output detailed markdown grouped by feature.\n\n"
            f"=== v1 FEATURES ===\n{_v1_feats or '(derive from the PRD)'}\n\n=== PRD (excerpt) ===\n{prd[:7000]}"),
            model=model_for("QA", model_map), max_tokens=8000)
        save_artifact(base_dir, name, "ACCEPTANCE.md", a)
        return a

    def _do_techspec():
        ts = call_agent("CTO", "", (
            "Produce a TECHNICAL SPEC in three clearly-headed sections:\n"
            "1) CONSTRAINTS & DECISIONS — target platform, supported devices/browsers, performance "
            "budgets, accessibility level, data/privacy/compliance, deployment target, and the "
            "libraries to USE vs AVOID — each with a one-line rationale.\n"
            "2) DATA MODEL — entities, fields, types and relationships.\n"
            "3) API CONTRACT — endpoints, methods and request/response shapes (omit if no backend).\n"
            "Be concrete. Output detailed markdown.\n\n"
            f"=== PRD (excerpt) ===\n{prd[:8000]}\n\n=== ENGINEERING GUIDELINES (excerpt) ===\n{tech_brief[:4000]}"),
            model=model_for("CTO", model_map), max_tokens=10000)
        save_artifact(base_dir, name, "TECH_SPEC.md", ts)
        if opt("datamodel"):
            dm = safe_json(call_agent("CTO", "Respond with valid JSON only.", (
                "From the project, produce a STRUCTURED data model and API contract. Return JSON only: "
                '{"entities": [{"name": str, "fields": [{"name": str, "type": str, "notes": str}]}], '
                '"endpoints": [{"method": str, "path": str, "request": str, "response": str, "notes": str}]}\n\n'
                f"=== PRD (excerpt) ===\n{prd[:7000]}\n\n=== TECH SPEC ===\n{ts[:6000]}"),
                model=model_for("CTO", model_map), json_mode=True, max_tokens=8000),
                default=_empty_datamodel())
            write_datamodel(base_dir, name, dm)
        return ts

    _waveB = []
    if opt("blueprint"):
        board["tasks"][6].update(state="dev", note="Mapping files & functions…")
        _waveB.append((6, _do_blueprint))
    else:
        board["tasks"][6].update(state="done", note="skipped")
    if opt("acceptance"):
        board["tasks"][7].update(state="dev", note="Writing acceptance criteria…")
        _waveB.append((7, _do_acceptance))
    else:
        board["tasks"][7].update(state="done", note="skipped")
    if opt("tech_spec"):
        board["tasks"][8].update(state="dev", note="Writing the technical spec…")
        _waveB.append((8, _do_techspec))
    else:
        board["tasks"][8].update(state="done", note="skipped")
    board["phase"] = "Blueprint, acceptance criteria & tech spec (in parallel)…"
    write_board(base_dir, name, board)
    rep(90, "Blueprint, acceptance criteria & tech spec drafting in parallel…")

    _resB = {}
    if _waveB:
        with ThreadPoolExecutor(max_workers=len(_waveB)) as _ex:
            _fB = {ti: _ex.submit(_stage, fn) for ti, fn in _waveB}
            for ti, f in _fB.items():
                _resB[ti] = f.result()
        for ti, _ in _waveB:
            board["tasks"][ti].update(state="done", note="Done ✓")
    acc = _resB.get(7, "") or ""
    techspec = _resB.get(8, "") or ""
    write_board(base_dir, name, board)

    # Red-team review + spec-readiness score (SKEPTIC) before development.
    if opt("redteam"):
        board["tasks"][9].update(state="dev", note="Red-teaming the plan…")
        board["phase"] = "Red-team review of the plan"
        write_board(base_dir, name, board)
        rep(97, "Skeptic red-teaming the plan…")
        rt = safe_json(call_agent("SKEPTIC", "Respond with valid JSON only.", (
            "Red-team this project BEFORE development. Review the PRD, plan, acceptance criteria and "
            "technical spec for missing requirements, contradictions, ambiguity, infeasibilities, "
            "security/privacy gaps, untestable criteria, and risky assumptions. Also SCORE the spec's "
            "readiness for an AI coding agent from 0-100 (clarity, completeness, testability). Return "
            "JSON only: {\"score\": 0-100, \"verdict\": str, \"issues\": [{\"severity\": "
            "\"high|med|low\", \"issue\": str, \"recommendation\": str}]}\n\n"
            f"=== PRD ===\n{prd[:7000]}\n\n=== PLAN ===\n{plan[:5000]}\n\n"
            f"=== ACCEPTANCE ===\n{acc[:5000]}\n\n=== TECH SPEC ===\n{techspec[:5000]}"),
            model=CEO_MODEL, max_tokens=8000), default={"score": None, "verdict": "", "issues": []})
        rt_lines = ["# Red-team review", "",
                    f"**Spec readiness:** {rt.get('score', '?')}/100 · **Verdict:** {rt.get('verdict', '')}", ""]
        for it in (rt.get("issues", []) or []):
            rt_lines.append(f"- **[{it.get('severity', '')}]** {it.get('issue', '')}")
            if it.get("recommendation"):
                rt_lines.append(f"  - ↳ {it.get('recommendation')}")
        save_artifact(base_dir, name, "REDTEAM.md", "\n".join(rt_lines))
        try:
            _ps = load_project_state(base_dir, name) or {}
            _ps["spec_score"] = rt.get("score")
            _ps["redteam_issues"] = len(rt.get("issues", []) or [])
            save_project_state(base_dir, name, _ps)
        except Exception:
            pass
        board["tasks"][9].update(state="done", note="Done ✓")
    else:
        board["tasks"][9].update(state="done", note="skipped")
    board["phase"] = "Documents ready"
    write_board(base_dir, name, board)

    rep(98, "Saving documents…")
    save_artifact(base_dir, name, "prd.md", prd)
    save_artifact(base_dir, name, "plan.md", plan)
    log_history(base_dir, name, "CEO", "Detailed multi-agent PRD and delivery plan generated.")
    log_history(base_dir, name, "CTO",
                "Authored engineering guidelines (architecture, clean code, quality standards).")
    rep(100, "Plan, PRD & engineering guidelines ready")
    return prd, plan


def start_planning(base_dir, name, seed, discussion, team, model_map, options=None):
    """Author the PRD/plan in a background daemon thread (progress in status.json)."""
    with _JOBS_LOCK:
        if _JOBS.get("plan:" + name) and _JOBS["plan:" + name].is_alive():
            return False
        t = threading.Thread(target=_run_planning,
                             args=(base_dir, name, seed, discussion, list(team), dict(model_map or {}),
                                   dict(options or {})),
                             daemon=True)
        _JOBS["plan:" + name] = t
        t.start()
        return True


def _run_planning(base_dir, name, seed, discussion, team, model_map, options=None):
    # Off the Streamlit thread — file IO + engine calls only, never st.*
    set_cost_context(base_dir, name, "planning")
    try:
        write_status(base_dir, name, phase="planning", running=True, progress=5,
                     action="Starting…", started_at=_now(), finished_at=None, error=None)

        def cb(pct, action):
            write_status(base_dir, name, phase="planning", progress=pct, action=action)

        prd, plan = build_plan_documents(base_dir, name, seed, discussion, team, model_map,
                                         on_progress=cb, options=options)
        # Fail-closed: a 'finished' planning run that produced nothing usable (empty PRD, or a PRD
        # dominated by a model-failure sentinel) must surface as an error — never as a green 'ready'
        # that the UI then loops on or shows as garbage. The error lands on the Retry/Back screen.
        if not (prd or "").strip() or is_engine_error(prd):
            write_status(base_dir, name, phase="planning", running=False, action="Planning failed",
                         error="Planning finished but didn't produce a usable PRD (the model calls "
                               "failed or returned nothing). Try again, or go back to add more detail.")
            return
        st = load_project_state(base_dir, name) or {}
        st["prd_content"] = prd
        st["plan_content"] = plan
        st["team"] = team
        st["stage"] = "planning"
        save_project_state(base_dir, name, st)
        write_status(base_dir, name, phase="planning", running=False, progress=100,
                     action="Plan & PRD ready", finished_at=_now())
    except Exception as e:
        write_status(base_dir, name, phase="planning", running=False, error=str(e),
                     action="Planning failed")
